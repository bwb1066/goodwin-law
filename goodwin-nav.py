from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

BASE = 'https://goodwinlaw.com'

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_plain_run(paragraph, text, bold=False, size=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

doc = Document()

# Title
title = doc.add_heading('Goodwin Law — Navigation Structure', 0)
title.runs[0].font.size = Pt(18)

# ── NAV DATA ──────────────────────────────────────────────────────────────────
# Format: (level, label, url_or_None, is_group_header)
# level 1 = top nav item, level 2 = sub-item with URL, level 3 = sub-item no URL

nav = [
    # EXPERTISE
    (1, 'Expertise', '/en/expertise'),

    (2, 'Litigation & Dispute Resolution', '/en/expertise#litigation'),
    (3, 'Antitrust & Competition', '/en/expertise/practices/antitrust-and-international-competition'),
    (3, 'Appellate and Supreme Court Litigation', '/en/expertise/practices/appellate-and-supreme-court-litigation'),
    (3, 'Business & Commercial Litigation', '/en/expertise/practices/business-litigation'),
    (3, 'Complex Litigation & Dispute Resolution', '/en/expertise/practices/complex-litigation-and-dispute-resolution'),
    (3, 'Data, Privacy & Cybersecurity', '/en/expertise/practices/data-and-privacy-and-cybersecurity'),
    (3, 'Employment', '/en/expertise/practices/employment'),
    (3, 'Environmental', '/en/expertise/practices/environmental'),
    (3, 'ERISA Litigation', '/en/expertise/practices/erisa-litigation'),
    (3, 'Financial Services Litigation', '/en/expertise/practices/financial-services-litigation'),
    (3, 'Government Contracts & Grants', '/en/expertise/practices/government-contracts-and-grants'),
    (3, 'Government Investigations, Enforcement & White Collar Defense', '/en/expertise/practices/enforcement-and-government-investigations'),
    (3, 'Intellectual Property Litigation', '/en/expertise/practices/intellectual-property-litigation'),
    (3, 'Investment Management Litigation', '/en/expertise/practices/investment-management-litigation'),
    (3, 'Life Sciences & Technology Arbitrations', '/en/expertise/practices/life-sciences-and-technology-arbitration-practice'),
    (3, 'Litigation', '/en/expertise/practices/litigation'),
    (3, 'National Security', '/en/expertise/practices/national-security'),
    (3, 'Products Litigation & Counseling', '/en/expertise/practices/products-litigation-and-counseling'),
    (3, 'Securities Litigation & SEC Enforcement', '/en/expertise/practices/securities-litigation-and-sec-enforcement'),
    (3, 'Trade Secrets, Employee Mobility & Non-Competes', '/en/expertise/practices/trade-secrets-employee-mobility-non-competes'),

    (2, 'Regulatory Compliance & Advisory', '/en/expertise#regulatory'),
    (3, 'Antitrust & Competition', '/en/expertise/practices/antitrust-and-international-competition'),
    (3, 'Banking', '/en/expertise/practices/banking'),
    (3, 'Data, Privacy & Cybersecurity', '/en/expertise/practices/data-and-privacy-and-cybersecurity'),
    (3, 'Digital Currency & Blockchain', '/en/expertise/practices/digital-currency-and-block-chain'),
    (3, 'Employment', '/en/expertise/practices/employment'),
    (3, 'Environmental', '/en/expertise/practices/environmental'),
    (3, 'ERISA & Executive Compensation', '/en/expertise/practices/erisa-and-executive-compensation'),
    (3, 'ESG & Impact', '/en/expertise/practices/esg-and-impact'),
    (3, 'Financial Restructuring', '/en/expertise/practices/financial-restructuring'),
    (3, 'Fintech', '/en/expertise/practices/fintech'),
    (3, 'Global Trade', '/en/expertise/practices/global-trade'),
    (3, 'Government Contracts & Grants', '/en/expertise/practices/government-contracts-and-grants'),
    (3, 'Insurtech', '/en/expertise/industries/insurtech'),
    (3, 'Investment Management', '/en/expertise/industries/investment-management'),
    (3, 'Late Stage Drug Development', '/en/expertise/practices/late-stage-drug-development'),
    (3, 'Life Sciences Regulatory & Compliance', '/en/expertise/practices/life-sciences-regulatory-and-compliance'),
    (3, 'National Security', '/en/expertise/practices/national-security'),
    (3, 'Private Investment Funds', '/en/expertise/industries/private-investment-funds'),
    (3, 'Proptech', '/en/expertise/industries/proptech'),
    (3, 'Public Company Advisory Practice', '/en/expertise/practices/public-company-advisory-practice'),
    (3, 'Risk Management & Insurance', '/en/expertise/practices/risk-management-and-insurance'),
    (3, 'Tax', '/en/expertise/practices/tax'),
    (3, 'Tax-Exempt Organizations', '/en/expertise/industries/tax-exempt-organizations'),
    (3, 'Trusts & Estates', '/en/expertise/practices/trusts-and-estate-planning'),

    (2, 'Transactions', '/en/expertise#transactions'),
    (3, 'Asset Management Transactions', '/en/expertise/practices/asset-management-transactions'),
    (3, 'Banking', '/en/expertise/practices/banking'),
    (3, 'Buyouts & Acquisitions', '/en/expertise/practices/buyouts-and-acquisitions'),
    (3, 'Capital Markets', '/en/expertise/practices/capital-markets'),
    (3, 'Corporate Social Responsibility Law & Policy', '/en/expertise/practices/corporate-social-responsibility-law-and-policy'),
    (3, 'Data, Privacy & Cybersecurity', '/en/expertise/practices/data-and-privacy-and-cybersecurity'),
    (3, 'Debt Finance', '/en/expertise/practices/debt-finance'),
    (3, 'Derivatives', '/en/expertise/practices/derivatives'),
    (3, 'Digital Currency & Blockchain', '/en/expertise/practices/digital-currency-and-block-chain'),
    (3, 'Employment', '/en/expertise/practices/employment'),
    (3, 'Environmental', '/en/expertise/practices/environmental'),
    (3, 'ERISA & Executive Compensation', '/en/expertise/practices/erisa-and-executive-compensation'),
    (3, 'ESG & Impact', '/en/expertise/practices/esg-and-impact'),
    (3, 'Financial Restructuring', '/en/expertise/practices/financial-restructuring'),
    (3, 'Government Contracts & Grants', '/en/expertise/practices/government-contracts-and-grants'),
    (3, 'Growth Equity', '/en/expertise/practices/growth-equity'),
    (3, 'Intellectual Property', '/en/expertise/practices/intellectual-property'),
    (3, 'Mergers & Acquisitions', '/en/expertise/practices/m-and-a-corporate-governance'),
    (3, 'New Company Formation', '/en/expertise/practices/new-company-formation'),
    (3, 'Private Investment Funds', '/en/expertise/industries/private-investment-funds'),
    (3, 'REITs & Real Estate M&A', '/en/expertise/industries/reits'),
    (3, 'Search Funds', '/en/expertise/practices/search-funds'),
    (3, 'Secondaries', '/en/expertise/practices/secondaries'),
    (3, 'Shareholder Activism and Takeover Defense', '/en/expertise/practices/shareholder-activism-and-takeover-defense'),
    (3, 'Special Purpose Acquisition Companies (SPACs)', '/en/expertise/practices/spacs'),
    (3, 'Tax', '/en/expertise/practices/tax'),
    (3, 'Tax-Exempt Organizations', '/en/expertise/industries/tax-exempt-organizations'),
    (3, 'Trusts & Estates', '/en/expertise/practices/trusts-and-estate-planning'),
    (3, 'Venture Capital', '/en/expertise/practices/venture-capital'),

    (2, 'Industries — Healthcare', '/en/expertise/industries/healthcare'),
    (3, 'Life Sciences', '/en/expertise/industries/life-sciences'),
    (3, 'Medtech', '/en/expertise/industries/medtech'),
    (3, "Women's Health and Wellness", '/en/expertise/industries/womens-health-and-wellness'),

    (2, 'Industries — Investment Funds', '/en/expertise/industries/investment-management'),
    (3, 'Private Investment Funds', '/en/expertise/industries/private-investment-funds'),
    (3, 'Registered Funds', '/en/expertise/practices/registered-funds'),
    (3, 'Hedge Funds', '/en/expertise/practices/hedge-funds'),

    (2, 'Industries — Life Sciences', '/en/expertise/industries/life-sciences'),
    (3, 'Medtech', '/en/expertise/industries/medtech'),

    (2, 'Industries — Private Equity', '/en/expertise/industries/private-equity'),
    (3, 'Buyouts & Acquisitions', '/en/expertise/practices/buyouts-and-acquisitions'),
    (3, 'Debt Finance', '/en/expertise/practices/debt-finance'),
    (3, 'Growth Equity', '/en/expertise/practices/growth-equity'),
    (3, 'Private Investment Funds', '/en/expertise/industries/private-investment-funds'),

    (2, 'Industries — Real Estate', '/en/expertise/industries/real-estate-industry'),
    (3, 'Hospitality & Leisure', '/en/expertise/industries/hospitality'),
    (3, 'Real Estate Finance & Restructuring', '/en/expertise/practices/real-estate-finance'),
    (3, 'Real Estate Investment Funds', '/en/expertise/practices/real-estate-private-investment-funds'),
    (3, 'Real Estate Joint Ventures', '/en/expertise/practices/real-estate-joint-ventures'),
    (3, 'REITs & Real Estate M&A', '/en/expertise/industries/reits'),

    (2, 'Industries — Technology', '/en/expertise/industries/technology-companies'),
    (3, 'Artificial Intelligence', '/en/expertise/practices/artificial-intelligence'),
    (3, 'Climate Tech', '/en/expertise/industries/climate-tech'),
    (3, 'Digital Currency & Blockchain', '/en/expertise/practices/digital-currency-and-block-chain'),
    (3, 'Industrial Technology & Advanced Manufacturing', '/en/expertise/industries/industrial-technology-and-advanced-manufacturing'),
    (3, 'Private Investment Funds', '/en/expertise/industries/private-investment-funds'),

    (2, 'Industries — Financial Services', '/en/expertise/industries/financial-services'),
    (3, 'Asset Management Transactions', '/en/expertise/practices/asset-management-transactions'),
    (3, 'Banking', '/en/expertise/practices/banking'),
    (3, 'Brokers, Exchanges, and Trading', '/en/expertise/practices/broker-dealer-exchange-and-trading'),
    (3, 'Consumer Financial Services', '/en/expertise/industries/consumer-financial-services'),
    (3, 'Fintech', '/en/expertise/practices/fintech'),
    (3, 'Insurance', '/en/expertise/industries/insurance'),
    (3, 'Insurtech', '/en/expertise/industries/insurtech'),
    (3, 'Investment Management', '/en/expertise/industries/investment-management'),

    # PEOPLE
    (1, 'People', '/en/people'),

    # CAREERS
    (1, 'Careers', '/en/careers-pages/careers'),
    (2, 'Legal Careers', '/en/careers-pages/legal-careers'),
    (2, 'Paralegal Careers', '/en/careers-pages/paralegal-careers'),
    (2, 'Science Advisor Careers', '/en/careers-pages/science-advisor-careers'),
    (2, 'Global Operations Team Careers', '/en/careers-pages/global-operations-careers'),

    # INSIGHTS
    (1, 'Insights & Resources', '/en/insights'),
    (2, 'Publications', '/en/insights?type=publications'),
    (2, 'Blog Posts', '/en/insights?type=blog'),
    (2, 'Newsletters', '/en/insights?type=newsletters'),
    (2, 'Case Studies', '/en/insights?type=case-studies'),

    # ABOUT
    (1, 'About Us', '/en/about/about-us'),
    (2, 'Our Story', '/en/about/about-us#our-story'),
    (2, 'The Goodwin Way', '/en/about/about-us#goodwin-way'),
    (2, 'Opportunity, Inclusion, and Belonging', '/en/about/opportunity-inclusion-belonging'),
    (2, 'Pro Bono', '/en/about/pro-bono'),
    (2, 'Firm Leadership', '/en/about/about-us#firm-leadership'),

    # LOCATIONS
    (1, 'Locations', '/en/locations'),
    (2, 'Boston', '/en/locations/boston'),
    (2, 'Los Angeles', '/en/locations/los-angeles'),
    (2, 'New York', '/en/locations/new-york'),
    (2, 'Orange County', '/en/locations/orange-county'),
    (2, 'Philadelphia', '/en/locations/philadelphia'),
    (2, 'San Francisco', '/en/locations/san-francisco'),
    (2, 'Santa Monica', '/en/locations/santa-monica'),
    (2, 'Silicon Valley', '/en/locations/silicon-valley'),
    (2, 'Washington DC', '/en/locations/washington-dc'),
    (2, 'Brussels', '/en/locations/brussels'),
    (2, 'Cambridge UK', '/en/locations/cambridge'),
    (2, 'London', '/en/locations/london'),
    (2, 'Luxembourg', '/en/locations/luxembourg'),
    (2, 'Munich', '/en/locations/munich'),
    (2, 'Paris', '/en/locations/paris'),
    (2, 'Hong Kong', '/en/locations/hong-kong'),
    (2, 'Singapore', '/en/locations/singapore'),

    # NEWS
    (1, 'News & Events', '/en/news-and-events'),

    # ALUMNI
    (1, 'Alumni', '/en/alumni'),
]

# ── RENDER ────────────────────────────────────────────────────────────────────

for level, label, path in nav:
    url = BASE + path

    if level == 1:
        doc.add_paragraph()  # spacer
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Pt(0)
        run = p.add_run()
        run.bold = True
        run.font.size = Pt(13)
        # Remove the run we just added and use hyperlink instead
        run._element.getparent().remove(run._element)
        add_hyperlink(p, label, url)
        # Make hyperlink bold via direct XML
        for r in p._p.findall('.//' + qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            b = OxmlElement('w:b')
            rPr.append(b)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '26')  # 13pt = 26 half-points
            rPr.append(sz)

    elif level == 2:
        p = doc.add_paragraph(style='List Bullet 2')
        add_hyperlink(p, label, url)
        for r in p._p.findall('.//' + qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '22')  # 11pt
            rPr.append(sz)

    elif level == 3:
        p = doc.add_paragraph(style='List Bullet 3')
        add_hyperlink(p, label, url)
        for r in p._p.findall('.//' + qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '20')  # 10pt
            rPr.append(sz)
            # grey color
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '555555')
            rPr.append(color)

out = '/Users/bwb1066/Documents/projects/goodwin-law/goodwin-nav.docx'
doc.save(out)
print(f'Saved: {out}')
