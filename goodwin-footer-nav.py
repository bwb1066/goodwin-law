from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

doc = Document()

title = doc.add_heading('Goodwin Law — Footer Navigation', 0)
title.runs[0].font.size = Pt(18)

footer_links = [
    ('Subscribe for Updates', 'https://www.goodwinlaw.com/en/subscribe'),
    ('Attorney Advertising', 'https://www.goodwinlaw.com/en/footer/attorney-advertising'),
    ('Cookie Policy', 'https://www.goodwinlaw.com/en/footer/cookie-policy'),
    ('Legal Notices', 'https://www.goodwinlaw.com/en/footer/legal-notices'),
    ('Privacy Policy', 'https://www.goodwinlaw.com/en/footer/privacy-policy'),
    ('Secure Login', 'https://www.goodwinlaw.com/en/footer/secure-login'),
    ('You@Goodwin', 'https://www.goodwinlaw.com/en/careers-pages/benefits#you-at-goodwin'),
]

social_links = [
    ('LinkedIn', 'https://www.linkedin.com/company/goodwin-law'),
    ('Instagram', 'https://www.instagram.com/join_goodwin/'),
    ('Twitter / X', 'https://twitter.com/goodwinlaw'),
    ('YouTube', 'https://www.youtube.com/@join_goodwin'),
]

doc.add_heading('Footer Links', level=1)

for label, url in footer_links:
    p = doc.add_paragraph(style='List Bullet')
    add_hyperlink(p, label, url)
    for r in p._p.findall('.//' + qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')
        rPr.append(sz)

doc.add_paragraph()
doc.add_heading('Social Media', level=1)

for label, url in social_links:
    p = doc.add_paragraph(style='List Bullet')
    add_hyperlink(p, label, url)
    for r in p._p.findall('.//' + qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')
        rPr.append(sz)

out = '/Users/bwb1066/Documents/projects/goodwin-law/goodwin-footer-nav.docx'
doc.save(out)
print(f'Saved: {out}')
